import os
import subprocess

from docx import Document
from rest_framework.exceptions import ValidationError

from root.settings import MEDIA_ROOT


def write_to_docx(text, id, name):
    file_name = f"{name[:40]}_{id}.docx"
    doc = Document()
    cleaned_output = (text.replace("***", "").replace("###", "").replace("##", "").replace("#", "")
                      .replace("**", "").replace("*", "").strip())
    doc.add_paragraph(cleaned_output)
    print(MEDIA_ROOT)
    if not os.path.exists(MEDIA_ROOT):
        os.makedirs(MEDIA_ROOT)
    file_path = os.path.join(MEDIA_ROOT, file_name)
    doc.save(file_path)
    return file_name


def write_to_docx2(text, id, name):
    file_name = f"{name[:40]}_{id}.docx"
    doc = Document()
    cleaned_output = (text.replace("***", "").replace("###", "").replace("##", "").replace("#", "")
                      .replace("**", "").replace("*", "").strip())
    doc.add_paragraph(cleaned_output)
    file_path = os.path.join(MEDIA_ROOT, file_name)
    doc.save(file_path)
    return file_path, file_name


def rewrite_docx(file_path, new_text):
    if not os.path.exists(file_path):
        raise ValidationError(f"The file {file_path} does not exist.", code="docx_file_not_found")

    doc = Document(file_path)

    for paragraph in doc.paragraphs:
        paragraph.clear()
    cleaned_output = new_text.replace("***", "").replace("###", "").replace("##", "").replace("#", "").replace("**",
                                                                                                               "").replace(
        "*", "").strip()
    doc.add_paragraph(cleaned_output)

    doc.save(file_path)

    return file_path


def read_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        raise ValidationError(detail="Fayl topilmadi", code="docx_file_not_found")


def convert_to_wav_v2(input_file_path, cpu_limit=50):
    try:
        output_file_path = os.path.splitext(input_file_path)[0] + ".wav"

        # Lower the process priority
        os.nice(10)

        # Command to use ffmpeg with cpulimit to convert the audio
        command = [
            'cpulimit', '-l', str(cpu_limit), '--',
            'ffmpeg', '-i', input_file_path,
            '-ar', '8000', '-ac', '1',
            output_file_path
        ]

        # Run the command
        process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        stdout, stderr = process.communicate()

        # Log the ffmpeg output
        stdout_log = stdout.decode()
        stderr_log = stderr.decode()

        print("FFmpeg stdout:\n", stdout_log)
        print("FFmpeg stderr:\n", stderr_log)

        if process.returncode != 0:
            raise Exception(f"FFmpeg error: {stderr_log}")

        # Remove the input file if conversion was successful
        if os.path.exists(input_file_path):
            os.remove(input_file_path)

        return output_file_path
    except Exception as e:
        raise e
