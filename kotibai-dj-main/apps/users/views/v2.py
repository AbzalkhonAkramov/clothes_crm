import os

import tiktoken
from docx import Document
from rest_framework import viewsets, status, generics
from rest_framework.exceptions import ValidationError
from rest_framework.pagination import LimitOffsetPagination
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.views import APIView

from root.settings import MEDIA_ROOT
from tasks import uzbek_voice, uzbek_voice_for_other, translate_tasks
from users.models import Project, Speechtotext, Summary, Article, Translate
from users.models.related import SpeechtotextForBot
from users.serializers.v2 import (
    ProjectSerializer, SpeechtotextSerializer,
    SummarySerializer, ArticleSerializer, TranslateSerializer, ProjectListSerializer, SpeechtotextForBotSerializer
)
from users.utils import lang_result, logger
from users.utils.openai_utils2 import send_request_openai
from users.utils.promt_data import summary_prompt, article_prompt, interview_prompt, news_prompt, reportage_prompt
from users.utils.speech_to_text import read_docx, rewrite_docx


class ProjectViewSet(viewsets.ModelViewSet):
    queryset = Project.objects.all()
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser]
    pagination_class = LimitOffsetPagination

    def get_queryset(self):
        return Project.objects.filter(user=self.request.user).order_by('-created_at')

    def get_serializer_class(self):
        if self.action == 'list':
            return ProjectListSerializer
        return ProjectSerializer


class BaseProjectRelatedViewSet(viewsets.ModelViewSet):
    permission_classes = [IsAuthenticated]

    def perform_create(self, serializer):
        project = Project.objects.get(id=self.kwargs['project_id'])
        serializer.save(user=self.request.user, project=project)

    def partial_update(self, request, *args, **kwargs):
        output_text = request.data.get('output_text', None)
        if output_text:
            queryset = self.get_object()
            queryset.output_text = output_text
            docx_file = queryset.output_docx.path
            rewrite_docx(file_path=docx_file, new_text=output_text)
            return Response(self.get_serializer(queryset.data), status=status.HTTP_200_OK)


class SpeechtotextCreateView(generics.RetrieveUpdateAPIView):
    queryset = Speechtotext.objects.all()
    serializer_class = SpeechtotextSerializer
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser]


class SummaryCreateView(BaseProjectRelatedViewSet):
    queryset = Summary.objects.all()
    serializer_class = SummarySerializer

    def create(self, request, *args, **kwargs):
        project = Project.objects.filter(id=kwargs.get('project_id')).first()
        if not project:
            raise ValidationError("Project does not exist.", code="project_does_not_exist")
        lang = request.data.get('lang')
        if lang not in ("uz-UZ", "ru-RU", "en-US", "ru-uz"):
            raise ValidationError("Invalid language: choose one: uz-UZ, ru-RU, en-US, ru-uz", code="invalid_lang")
        stt_obj = getattr(project, 'stt', None)
        if stt_obj and project.stt.output_docx or project.action_type == project.ActionChoices.STT:
            if not project.stt.output_docx:
                raise ValidationError("STT output text is empty.", code="stt_empty")
            result_text = read_docx(project.stt.output_docx.path)
        else:
            if project.input_text:
                result_text = project.input_text
            elif project.input_file:
                result_text = read_docx(project.input_file.path)
            else:
                raise ValidationError("Project Input text or input file is empty.", code="project_input_empty")

        encoding = tiktoken.encoding_for_model("gpt-3.5-turbo")
        tokens = encoding.encode(result_text)
        num_tokens = len(tokens)
        if num_tokens > 128000:
            tokens = tokens[:127800]
            result_text = encoding.decode(tokens)
        model = "gpt-4o"
        result_answer = send_request_openai(summary_prompt, result_text, model, project.user)
        print(result_answer, "\n\n++++++++++++++++++++\n\n++++++++++++++++++++", num_tokens)
        tr_lang_prompt = lang_result(f"auto-{lang[:2]}")
        result_answer = send_request_openai(prompt_system=tr_lang_prompt, prompt_user=result_answer, model="gpt-4o",
                                            user=project.user)

        print(result_answer, "\n\n++++++++++++++++++++\n\n++++++++++++++++++++", num_tokens)
        doc = Document()
        doc.add_paragraph(result_answer)
        translated_path = f'{MEDIA_ROOT}summary_{lang}_{project.name}.docx'
        doc.save(translated_path)
        translated_file_name = os.path.basename(translated_path)
        summery = Summary.objects.create(user=request.user, project=project, lang=lang, output_text=result_answer,
                                         output_docx=translated_file_name)
        response_data = SummarySerializer(summery, context={'request': request}).data
        return Response(response_data, status=status.HTTP_201_CREATED)


class ArticleCreateView(BaseProjectRelatedViewSet):
    queryset = Article.objects.all()
    serializer_class = ArticleSerializer

    def create(self, request, *args, **kwargs):
        project = Project.objects.filter(id=kwargs.get('project_id')).first()
        if not project:
            raise ValidationError("Project does not exist.", code="project_does_not_exist")
        lang = request.data.get('lang')
        if lang not in ("uz-UZ", "ru-RU", "en-US", "ru-uz"):
            raise ValidationError("Invalid language: choose one: uz-UZ, ru-RU, en-US, ru-uz", code="invalid_lang")
        article_type = request.data.get('type')

        if article_type == "article":
            prompt = article_prompt
        elif article_type == "interview":
            prompt = interview_prompt
        elif article_type == "news":
            prompt = news_prompt
        elif article_type == "reportage":
            prompt = reportage_prompt
        else:
            raise ValidationError('Invalid article type', code="invalid_article_type")

        stt_obj = getattr(project, 'stt', None)
        if stt_obj and project.stt.output_docx or project.action_type == project.ActionChoices.STT:
            if not project.stt.output_docx:
                raise ValidationError("STT output text is empty.", code="stt_empty")
            result_text = read_docx(project.stt.output_docx.path)
        else:
            if project.input_text:
                result_text = project.input_text
            elif project.input_file:
                result_text = read_docx(project.input_file.path)
            else:
                raise ValidationError("Project Input text or input file is empty.", code="project_input_empty")
        encoding = tiktoken.encoding_for_model("gpt-3.5-turbo")
        tokens = encoding.encode(result_text)
        num_tokens = len(tokens)
        model = "gpt-4o"
        if num_tokens > 128000:
            tokens = tokens[:127700]
            result_text = encoding.decode(tokens)

        result_answer = send_request_openai(prompt, result_text, model, project.user)
        print(prompt, "prompt ----------")
        print(result_answer, "\n\n++++++++++++++++++++\n\n++++++++++++++++++++", num_tokens)
        tr_lang_prompt = lang_result(f"auto-{lang[:2]}")
        result_answer = send_request_openai(prompt_system=tr_lang_prompt, prompt_user=result_answer, model="gpt-4o",
                                            user=project.user)
        doc = Document()
        doc.add_paragraph(result_answer)
        translated_path = f'{MEDIA_ROOT}article_{lang}_{project.name}.docx'
        doc.save(translated_path)
        translated_file_name = os.path.basename(translated_path)
        article = Article.objects.create(user=request.user, project=project, lang=lang, type=article_type,
                                         output_text=result_answer, output_docx=translated_file_name)
        response_data = ArticleSerializer(article, context={'request': request}).data
        return Response(response_data, status=status.HTTP_201_CREATED)


class TranslateCreateView(BaseProjectRelatedViewSet):
    queryset = Translate.objects.all()
    serializer_class = TranslateSerializer
    permission_classes = [IsAuthenticated]

    def create(self, request, *args, **kwargs):
        project = Project.objects.filter(id=kwargs.get('project_id')).first()
        user = self.request.user
        if not project:
            raise ValidationError("Project does not exist.", code="project_does_not_exist")
        lang = request.data.get('lang')
        if lang not in ("uz-UZ", "ru-RU", "en-US", "ru-uz"):
            raise ValidationError("Invalid language: choose one: uz-UZ, ru-RU, en-US, ru-uz", code="invalid_lang")
        stt_obj = getattr(project, 'stt', None)
        if stt_obj and project.stt.output_docx or project.action_type == project.ActionChoices.STT:
            if not project.stt.output_docx:
                raise ValidationError("STT output text is empty.", code="stt_empty")
            result_text = read_docx(project.stt.output_docx.path)
        else:
            if project.input_text:
                result_text = project.input_text
            elif project.input_file:
                result_text = read_docx(project.input_file.path)
            else:
                raise ValidationError("Project Input text or input file is empty.", code="project_input_empty")

        tr_lang_prompt = lang_result(f"auto-{lang[:2]}")
        translate_tasks.delay(project.id, user.id, result_text, lang, tr_lang_prompt)
        project.loading = True
        project.save()

        return Response({"detail: keyinroq javob keladi"}, status=status.HTTP_201_CREATED)


class SpeechtotextForBotCreateApiView(generics.CreateAPIView):
    queryset = SpeechtotextForBot.objects.all()
    serializer_class = SpeechtotextForBotSerializer

    def perform_create(self, serializer):
        instance = serializer.save()
        logger.info("New instance created: {}".format(instance))


class GetUzbekVoiceText(APIView):
    def post(self, request):
        voice_id = request.data.get('id')
        text = request.data.get('result')['text']
        print("get uzbek voice, +++++++++++++++++++++++++++++++++++++++++")
        uzbek_voice.delay(voice_id, text)

        return Response(request.data)


class GetUzbekVoiceForOther(APIView):
    def post(self, request):
        voice_id = request.data.get('id', None)
        text = request.data.get('result', None)['text']
        print(request.data, "++++++++++++++++++++++++\n++++++++++++++++++\n++++++++++++++")
        uzbek_voice_for_other.delay(voice_id, text)
        return Response(request.data)
